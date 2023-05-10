import mysql.connector
import os
import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
import numpy as np
from matplotlib.ticker import MaxNLocator
import matplotlib
import math
from sklearn.cluster import KMeans
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Mm
from docx.shared import Inches
import io
import base64
from docxcompose.composer import Composer
import scipy.stats as st
from pyecharts.charts import Geo
from pyecharts.charts import Map
from sklearn.metrics import r2_score
import pandas as pd
import seaborn as sns
from matplotlib.ticker import MaxNLocator
from pyecharts import options as opts
from pyecharts.charts import Line,Grid,Bar,Pie,Radar
from pyecharts.faker import Faker
from pyecharts.globals import ThemeType
from pyecharts.globals import ChartType
import json
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pyecharts.render import make_snapshot
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from snapshot_selenium import snapshot
import yaml
import subprocess

class mydb:#数据库操作类

    def __init__(self,host,user,password,database,port,charset):
        self.host=host
        self.user=user
        self.password=password
        self.database=database
        self.port=port
        self.charset=charset

    def myconnect(self):#连接数据库
        conn = mysql.connector.connect(
            host=self.host,
            user=self.user,
            password=self.password,
            database=self.database,
            port=self.port,
            charset = self.charset
            # auth_plugin='mysql_native_password'
        )
        return conn

    def myexecu(self,execuword):#执行查询语句并保存在data中
        self.conn=self.myconnect()
        cursor=self.conn.cursor()
        cursor.execute(execuword)
        self.data = cursor.fetchall()
        self.conn.commit()
        cursor.close()


    def endconn(self):#断开数据库连接
        self.conn.close()

class cladata:#基类

    def __init__(self,db,title):#将data内容传递给cladata类
        self.data=db.data
        self.title=title
        self.db=db


class authoranaly(cladata):#申请人分析

    def authorpie(self,dic,title):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        mydic={}
        result=[]
        n=0
        for key,value in dic.items():
            n+=1
            if n<=20:
                mydic[key]=value
            else:
                break
        for key,value in mydic.items():
            result.append([key,value])
        c.add("",result)
        c.set_global_opts(title_opts=opts.TitleOpts(title="专利数量前20位申请人分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
        return c.dump_options_with_quotes()


    def authorbar(self,dic):#绘制柱状图
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        X, Y = [], []
        n=0
        for key,value in dic.items():
            n+=1
            if n<=20:
                X.append(key)
                Y.append(int(value))
            else:
                break
        c.add_xaxis(X)
        c.add_yaxis(series_name="申请人", y_axis=Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title="专利数量前20位申请人分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        return c.dump_options_with_quotes()

    def authorlist(self):#将数据中所有作者解析
        dic={}
        for i in self.data:
            mystr=i[1]
            list=str(mystr).split(";")
            for h in list:
                if h[0]==' ':
                    h=h[1:]
                dic[h]=0
        return dic

    def listsort(self,dic,id):#给字典赋值并排序
        total=0
        for i in dic.keys():
            i=str(i)
            self.db.myexecu("select count(inventor_list) as cnt from( select inventor_list from patent.patent where  (inventor_list like'%; "+i+"'"+" or inventor_list like'%"+i+";%'"+" or inventor_list like '%;"+i+"'"+" or inventor_list ="+"'"+i+"'"+") and id in"+id+" ) as ccc")
            num=str(self.db.data)
            num=num[2:-3]
            num=int(num)
            dic[i]=num
        for i in dic.values():
            total+=i
        self.total=total
        a1=sorted(dic.items(),key=lambda x:x[1],reverse = True)
        dic=dict(a1)
        self.dic=dic
        return dic


class areaanaly(cladata):#地域分析


    def certainarea(self,name,id):#某个省份的信息
        self.db.myexecu("SELECT  * FROM patent.patent where id in "+id+" and ( substring(applicant_area,1,INSTR(applicant_area,'(')-1)='"+name+"' or substring(applicant_area,1,INSTR(applicant_area,';')-1)='"+name+"')")
        return self.db.data

    def newdic(self):#创建一个字典
        dic={}
        for i in self.data:
            dic[i[1]]=i[0]
        return dic

    def areabar(self,dic):#绘制柱状图
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        X, Y = [], []
        for key,value in dic.items():
                X.append(key)
                Y.append(int(value))
        c.add_xaxis(X)
        c.add_yaxis(series_name="地区", y_axis=Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title="地域分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        return c.dump_options_with_quotes()

    def areapie(self,dic):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        result=[]
        for key,value in dic.items():
            result.append([key,value])
        c.add("",result)
        c.set_global_opts(title_opts=opts.TitleOpts(title="地域分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
        return c.dump_options_with_quotes()


class trendanaly(cladata):#趋势分析



    def opendic(self,stime,etime):#每遇到一个新类型就新建一个新字典
        dic={}
        i=stime
        while i<=etime:
            dic[i]=0
            i+=1
        return dic

    def newdic(self,stime,etime):#创建值为字典的字典对
        updic={}
        for h in self.data:
            updic[h[0]]=self.opendic(stime,etime)
        for h in self.data:
            updic[h[0]][float(h[1])]=h[2]
        return updic

    def trendpie(self,dic):#绘制每个类别的饼状图
        output=[]
        for key,value in dic.items():
            output.append(self.pietrend(value,key))
        return output


    def pietrend(self,dic,title):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        result=[]

        for key,value in dic.items():
            if value != 0:
                result.append([str(key)+"年",value,])
        c.add("",result)
        c.set_global_opts(title_opts=opts.TitleOpts(title="分类号"+str(title)+"趋势分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
        return c.dump_options_with_quotes()

    def trendline(self,dic,stime,etime):#绘制折线图
        mytimelist=[]
        mytime=stime
        while mytime<=etime:
            mytimelist.append(str(mytime)+"年")
            mytime+=1
        c =Line(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        c.add_xaxis(mytimelist)
        for key,value in dic.items():
            Y=[]
            for upvalue in value.values():
                Y.append(upvalue)
            c.add_yaxis( series_name=key,y_axis=Y, is_connect_nones=True)
            print(Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title=str(stime)+"年至"+str(etime)+"年趋势分析折线图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        return c.dump_options_with_quotes()


    def trendbar(self,dic,stime,etime):#绘制时间趋势柱状图
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        mytimelist=[]
        mytime=stime
        while mytime<=etime:
            mytimelist.append(str(mytime)+"年")
            mytime+=1
        c.add_xaxis(mytimelist)
        mydic={}
        for key,value in dic.items():
            for uvalue in value.values():
                if uvalue !=0:
                    mydic[key]=value
                    break
        for upkey,upvalue in mydic.items():
            Y = []
            for value in upvalue.values():
                Y.append(int(value))
            c.add_yaxis( series_name=upkey,y_axis=Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title=str(stime)+"年至"+str(etime)+"年趋势分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        return c.dump_options_with_quotes()

class pdfcreator:#pdf生成器

    def __init__(self,id,db):
        self.id=id
        self.db=db
        self.db.myexecu("SELECT item_type,corr_id FROM patent.report_content_item where report_id="+str(id))
        self.data=self.db.data
        yaml_path= os.path.dirname(os.path.dirname(__file__)) + "/config/config.yml"
        with open(yaml_path,"r",encoding="utf-8") as f:
            data=yaml.load(f,Loader=yaml.FullLoader)
        self.path=data['pdf-output-path']
        self.searchnum=0
        self.newnum=0

    def novelty_table_design(self, document, related_dic):
        # 新颖性比对分析，相关权利要求的信息使用表格绘制
        for relate_sig_id in range(len(related_dic["related"])):
            # "related":[{"signory_text","patent_title","suggestion"}]

            table = document.add_table(rows=2, cols=3, style='TableGrid')
            for row in range(0, 4):
                table.add_row()

            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            # table.cell(0, 0).merge(table.cell(0, 1))
            # table.cell(0, 0).text = "相关权利要求"+str(relate_sig_id+1)
            # table.cell(0, 2).text = related_dic["related"][relate_sig_id]["signory_text"]

            # 0行和1行
            table.cell(0, 1).text = "相关性"
            # TODO: 相似度得分需要从related_dic或者哪里拿一下
            table.cell(1, 1).text = related_dic["related"][relate_sig_id]["score"]
            table.cell(0, 0).merge(table.cell(1, 0))
            table.cell(0, 0).text = "相关权利要求" + str(relate_sig_id + 1)
            table.cell(0, 2).merge(table.cell(1, 2))
            table.cell(0, 2).text = related_dic["related"][relate_sig_id]["signory_text"]

            # 2行和3行，所属专利信息
            table.cell(2, 1).text = "标题"
            table.cell(2, 2).text = related_dic["related"][relate_sig_id]["patent_title"]
            table.cell(3, 1).text = "申请号"
            # TODO：对于检索到的权利要求所属申请号，如果拿不到就删掉这一行
            table.cell(3, 2).text = related_dic["related"][relate_sig_id]["patent_code"]

            table.cell(2, 0).merge(table.cell(3, 0))
            table.cell(2, 0).text = "所属专利信息"

            table.cell(4, 0).merge(table.cell(4, 1))
            table.cell(4, 0).merge(table.cell(4, 2))
            table.cell(4, 0).text = "审查意见"
            table.cell(5, 0).merge(table.cell(5, 1))
            table.cell(5, 0).merge(table.cell(5, 2))
            table.cell(5,0).text = self.novelty_string_change(related_dic["related"][relate_sig_id]["suggestion"])

            for row in range(5):
                for column in range(3):
                    if column == 0 or column == 1:
                        table.cell(row, column).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.cell(row, column).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(0.9)
            table.columns[2].width = Inches(4)

            p = document.add_paragraph()

    def novelty_string_change(self,novelty_string):
        # TODO：如果需要改一下审查意见的格式，直接到这里来改就行，预留一个函数
        # novel_paragraph = "采取相关关系词对、新颖性判断规则对待分析权利要求和相关权利要求进行分析后，获得比对结果为：可能破坏带申请专利技术特征的依据点共有{}条。具体分析如下。\n".format(novelty_string.count("将破坏"))
        # 设想是修改一下(i)相关关系词对(ii)规则判断，规则判断的每一个小点都用一个·或者一个-代替
        novelty_string_new = novelty_string.replace("相关关系词对：","(i)相关关系词对:")
        novelty_string_new = novelty_string_new.replace("规则判断：", "(ii)规则判断:")
        # novelty_string_new = novelty_string
        return novelty_string_new


    def finalword(self):#生成最终的word
        jiannum=0
        snum=[]
        nnum=[]
        numa=0
        a=1
        pagenum={1:"一、",2:"二、",3:"三、"}
        document = Document()
        mei=document.sections[0].header.paragraphs[0]
        m=document.sections[0].header
        text=mei.add_run("专利分析报告")
        text.font.size = Pt(10)
        text.bold = True
        text.font.name = 'Arial'
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        mei.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 修改页眉，85-8
        # line = mei.add_run("_" * 77)
        # line.font.size = Pt(10)

        # 通过分区添加页眉横线
        m.add_paragraph("\t\t").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para = m.paragraphs[-1]
        # style1:与页边距齐平，这样对不齐
        # header_para.paragraph_format.left_indent = 0
        # header_para.paragraph_format.right_indent = 0
        # style2:横穿整个页面故段落边距紧贴页边，通过调整数值调整横线和页边的距离
        header_para.paragraph_format.left_indent = - document.sections[0].left_margin  + docx.shared.Inches(0.8)
        header_para.paragraph_format.right_indent = - document.sections[0].right_margin + docx.shared.Inches(2.0)
        # 设置制表位分区
        tab_stops = header_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(docx.shared.Inches(3.5), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(docx.shared.Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        # 设置横线样式
        header_para.runs[-1].underline = WD_UNDERLINE.THICK  # 厚 DOUBLE双下划线
        
        p=document.add_paragraph("专利分析报告")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(27)
        p.runs[0].font.name = '宋体'
        p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p=document.add_paragraph("目录")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.name = '宋体'
        p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if(len(self.addsearchresult())!=0):
            pp=document.add_paragraph(pagenum[a]+"专利检索结果及分析")
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.runs[0].font.size = Pt(12)
            pp.runs[0].font.name = '宋体'
            pp.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            a+=1
        if(len(self.addnovelresult())!=0):
            # 修改标题名称
            aa=document.add_paragraph(pagenum[a]+"新颖性创造性比对结果及分析")
            aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
            aa.runs[0].font.size = Pt(12)
            aa.runs[0].font.name = '宋体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            a+=1
        bb=document.add_paragraph(pagenum[a]+"小结")
        bb.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bb.runs[0].font.size = Pt(12)
        bb.runs[0].font.name = '宋体'
        bb.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        a+=1
        document.paragraphs[a].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        ji=1
        npagenum={1:"一、",2:"二、",3:"三、"}
        if(len(self.addsearchresult())!=0):
            for i in self.data:
                if i[0]=="检索结果":
                    snum.append(i[1])
            aa=document.add_paragraph(npagenum[ji]+"专利检索结果及分析")
            aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
            aa.runs[0].font.size = Pt(14)
            aa.runs[0].font.name = '宋体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            for a in self.addsearchresult():
                self.searchnum+=1
                aa=document.add_paragraph(str(self.searchnum)+"、检索结果集"+str(self.searchnum)+":")
                aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                aa.runs[0].font.size = Pt(13)
                aa.runs[0].font.name = '宋体'
                aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                aa=document.add_paragraph("(1)检索结果内容")
                aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                aa.runs[0].font.size = Pt(12)
                aa.runs[0].font.name = '宋体'
                aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                table = document.add_table(rows=1, cols=3,style='TableGrid')
                rows = table.rows[0]
                for cell in rows.cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                # 修改bug:检索结果表格表头不在第一行
                table.rows[0].cells[0].text = "标题"
                table.rows[0].cells[1].text = "申请人列表"
                table.rows[0].cells[2].text = "申请时间"
                for b in a:
                    for c in b:
                        jiannum+=1
                        document.styles['Normal'].font.name = u'宋体'
                        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        cells = table.add_row().cells
                        cells[0].text=c[0]
                        cells[1].text=c[1]
                        cells[2].text=c[2]
                document.add_paragraph(" ")
                if len(self.addpicresult(snum[self.searchnum-1]))!=0:
                    aa=document.add_paragraph("(2)统计分析结果")
                    aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    aa.runs[0].font.size = Pt(12)
                    aa.runs[0].font.name = '宋体'
                    aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    for i in self.addpicresult(snum[self.searchnum-1]):
                        if i["type"]=="柱状":
                            c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add_xaxis(i["xaxis"])
                            if len(i["data"])==1:
                                (key, value), = i["data"].items()
                                c.add_yaxis( series_name=key,y_axis=value)
                            else:
                                for key,value in i["data"].items():
                                    c.add_yaxis( series_name=key,y_axis=value)
                            c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"bar.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run("")
                            inline_shape=run.add_picture("bar.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
                            run.add_picture(self.path+str(self.id)+"bar.png",width=Inches(7.0))
                        if i["type"]=="折线":
                            c =Line(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add_xaxis(i["xaxis"])
                            if len(i["data"])==1:
                                (key, value), = i["data"].items()
                                c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                            else:
                                for key,value in i["data"].items():
                                    c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                            c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"line.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run("")
                            inline_shape=run.add_picture(self.path+str(self.id)+"line.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
                        if i["type"]=="饼状":
                            c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add("",i["data"])
                            c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"pie.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run("")
                            inline_shape=run.add_picture(self.path+str(self.id)+"pie.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
            ji+=1
            document.add_paragraph(" ")
        if(len(self.addnovelresult())!=0):
            for i in self.data:
                if i[0]=="新颖性比对结果":
                    nnum.append(i[1])
            numa=1
            aa=document.add_paragraph(npagenum[ji]+"新颖性创造性比对结果及分析")
            aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
            aa.runs[0].font.size = Pt(14)
            aa.runs[0].font.name = '宋体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            for related_dic in self.addnovelresult():
                aa=document.add_paragraph(str(numa)+"、新颖性比对结果"+str(numa)+":")
                aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                aa.runs[0].font.size = Pt(13)
                aa.runs[0].font.name = '宋体'
                aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                # 逻辑修改：如果要呈现比对结果和图表分析结果，则加上（1）(2)；如果没有图表，小括号级的标题不加
                if (len(self.addnewpicresult(nnum[numa - 2])) != 0):
                    content_title = "(1)新颖性比对结果内容"
                    aa = document.add_paragraph(content_title)
                    aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    aa.runs[0].font.size = Pt(12)
                    aa.runs[0].font.name = '宋体'
                    aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                # 逻辑：首先获取原权利要求并展示，其次讲相关权利要求信息写入表格并进行展示
                p = document.add_paragraph()
                p.style.font.name = '宋体'
                p.bold = True
                p.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                p.add_run("原权利要求：\n")
                p.add_run("\t"+related_dic["origin_sig_text"] + "\n")
                p = document.add_paragraph()
                p.add_run("根据原权利要求进行检索、比对后，得到相关权利要求如下：")
                self.novelty_table_design(document, related_dic)
                if(len(self.addnewpicresult(nnum[numa-2]))!=0):
                    aa = document.add_paragraph("(2)新颖性分析结果统计")
                    aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    aa.runs[0].font.size = Pt(12)
                    aa.runs[0].font.name = '宋体'
                    aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                for i in self.addnewpicresult((nnum[numa-2])):
                        if i["type"]=="柱状":
                            c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add_xaxis(i["xaxis"])
                            if len(i["data"])==1:
                                (key, value), = i["data"].items()
                                c.add_yaxis( series_name=key,y_axis=value,bar_width="40%")
                            else:
                                for key,value in i["data"].items():
                                    c.add_yaxis( series_name=key,y_axis=value)
                            c.set_global_opts(xaxis_opts=opts.AxisOpts(name = "权利要求编号",splitline_opts=opts.SplitLineOpts(is_show=False)),yaxis_opts=opts.AxisOpts(name="数量",splitline_opts=opts.SplitLineOpts(is_show=False)),title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"bar.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            run = paragraph.add_run("")
                            inline_shape= run.add_picture(self.path+str(self.id)+"bar.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
                            document.add_paragraph("   ")
                        if i["type"]=="折线":
                            c =Line(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add_xaxis(i["xaxis"])
                            if len(i["data"])==1:
                                (key, value), = i["data"].items()
                                c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                            else:
                                for key,value in i["data"].items():
                                    c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                            c.set_global_opts(xaxis_opts=opts.AxisOpts(name = "权利要求编号",splitline_opts=opts.SplitLineOpts(is_show=False)),title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"line.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run("")
                            inline_shape=run.add_picture(self.path+str(self.id)+"line.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
                            document.add_paragraph("   ")
                        if i["type"]=="饼状":
                            c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                            c.add("",i["data"])
                            c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                            c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
                            make_snapshot(snapshot, c.render(), self.path+str(self.id)+"pie.png",is_remove_html=True)
                            paragraph=document.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run("")
                            inline_shape=run.add_picture(self.path+str(self.id)+"pie.png",width=Inches(7.0))
                            inline_shape.height = Cm(8.06)
                            inline_shape.width = Cm(14.5)
                            document.add_paragraph("   ")

        ji+=1
        document.add_paragraph(" ")
        aa=document.add_paragraph(npagenum[ji]+"小结")
        aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
        aa.runs[0].font.size = Pt(14)
        aa.runs[0].font.name = '宋体'
        aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p=document.add_paragraph()
        p.style.font.name = '宋体'
        p.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        result=[]
        result.append(self.searchnum)
        result.append(jiannum)
        result.append(numa)
        idstr="("
        for i in nnum:
            idstr+=str(i)+","
        idstr=idstr[0:-1]
        idstr+=")"
        if len(nnum)==0:
            p.add_run("\t在专利预评估分析过程中,得到了"+str(result[0])+"篇检索结果集,其中共包括"+str(result[1])+"篇相关专利。对"+str(result[2])+"条权利要求进⾏了新颖性分析,在与各自相关权利要求的⽐较中,探测相关词0对,新颖性评判规则相关点0条,其中涉及上下位概念相关点0条,惯⽤⼿段直接置换0条,数字范围相关点0条,新颖性⻛险点0条。")
        else:
            self.db.myexecu("SELECT sum(word_pairs_sum),sum(trigger_rules_sum),sum(hyponym_hypernym_sum),sum(direct_substitution_sum),sum(numeric_range_sum),sum(destroy_sum)FROM patent.novelty_ana_result where novelty_ana_id in "+idstr)
            p.add_run("在专利预评估分析过程中,得到了"+str(result[0])+"篇检索结果集,其中共包括"+str(result[1])+"篇相关专利。对"+str(result[2])+"条权利要求进⾏了新颖性分析,在与各自相关权利要求的⽐较中,探测相关词"+str(self.db.data[0][0])+"对,新颖性评判规则相关点"+str(self.db.data[0][1])+"条,其中涉及上下位概念相关点"+str(self.db.data[0][2])+"条,惯⽤⼿段直接置换"+str(self.db.data[0][3])+"条,数字范围相关点"+str(self.db.data[0][4])+"条,新颖性⻛险点"+str(self.db.data[0][5])+"条。")
        self.db.myexecu("SELECT report_id,report_name from patent.report2generate where report_id= "+str(self.id))
        self.name=self.db.data[0][1]
        document.save(self.path+str(self.name)+".docx")

    def addsearchresult(self):#专利检索内容增加
        tabledata=[]
        mydata=[]
        self.searchid=[]
        for i in self.data:
            if i[0]=="检索结果":
                updata=[]
                self.db.myexecu("SELECT patent_id FROM patent.analysis_collection_item where collection_id="+str(i[1]))
                for a in self.db.data:
                    mydata.append(a[0])
                for a in mydata:
                    uppdata=[]
                    self.db.myexecu("SELECT title,inventor_list,application_date FROM patent.patent where id ="+str(a))
                    for b in self.db.data:
                        uppdata.append(b)
                    updata.append(uppdata)
                tabledata.append(updata)
        return tabledata

    def addnovelresult(self):  # 新颖性内容增加
        # 改后：mydata存放的对象为一个新颖性比对结果的全部
        # updata格式： {”origin_sig_text“:str,"related":[{"signory_text","patent_title","suggestion"}]}
        mydata = []
        for i in self.data:
            if i[0] == "新颖性比对结果":
                signories = {}
                self.db.myexecu("SELECT ori_signory FROM patent.novelty_ana_result where novelty_ana_id=" + str(i[1]))
                signories["origin_sig_text"] = str(self.db.data[0][0])
                signories["related"] = []
                self.db.myexecu(
                    "SELECT relevant_sig,compare_result,ori_patent_title, score, patent_code FROM patent.novelty_ana_item where novelty_ana_id=" + str(
                        i[1]))
                # 这里获取了新颖性的主权项并存放成词典
                for a in self.db.data:
                    temp_dic = {"signory_text": a[0], "patent_title": a[2], "suggestion": a[1], "score": a[3], "patent_code": a[4]}
                    signories["related"].append(temp_dic)
                mydata.append(signories)
        return mydata

    def addpicresult(self,id):#图片内容
        outresult=[]
        tabledata=[]
        self.db.myexecu("SELECT stats_res_id FROM patent.stats_ana_result where analysis_collection_id="+str(id))
        mystr=str(self.db.data)
        mystr=mystr[2:-3]
        self.db.myexecu("SELECT option_json FROM patent.stats_ana_item where stats_ana_id="+str(mystr))
        for i in self.db.data:
            myjson=json.loads(str(i[0]))
            mytitle=myjson["title"][0]["text"]
            type=mytitle[-3:-1]
            if type=="柱状":
                barresult={}
                barresult["type"]="柱状"
                barresult["title"]=mytitle
                time=[]
                for a in myjson["xAxis"][0]["data"]:
                    time.append(a)
                barresult["xaxis"]=time
                mydata={}
                for a in myjson["series"]:
                    mydata[a["name"]]=a["data"]
                barresult["data"]=mydata
                outresult.append(barresult)
            if type=="折线":
                barresult={}
                barresult["type"]="折线"
                barresult["title"]=mytitle
                time=[]
                for a in myjson["xAxis"][0]["data"]:
                    time.append(a)
                barresult["xaxis"]=time
                mydata={}
                for a in myjson["series"]:
                    updata=[]
                    for b in a["data"]:
                        updata.append(b[1])
                    mydata[a["name"]]=updata
                barresult["data"]=mydata
                outresult.append(barresult)
            if type=="饼状":
                barresult={}
                barresult["type"]="饼状"
                barresult["title"]=mytitle
                time=[]
                mydata=[]
                a = myjson["series"][0]["data"]
                for b in a:
                    mydata.append([b["name"],b["value"]])
                barresult["data"]=mydata
                outresult.append(barresult)
        return outresult



    def addnewpicresult(self,id):#图片内容
        outresult=[]
        tabledata=[]
        self.db.myexecu("SELECT id FROM patent.novelty_stats_result where novelty_ana_result_id="+str(id))
        mystr=str(self.db.data)
        mystr=mystr[2:-3]
        self.db.myexecu("SELECT option_json FROM patent.novelty_stats_item where novelty_stats_id="+mystr)
        for i in self.db.data:
            myjson=json.loads(str(i[0]))
            mytitle=myjson["title"][0]["text"]
            if mytitle=="各比对结果相关词和规则相关点数量" or mytitle=="规则相关点各类型数量" or mytitle=="各比对结果疑似新颖性风险点":
                barresult={}
                barresult["type"]="柱状"
                barresult["title"]=mytitle
                time=[]
                for a in myjson["xAxis"][0]["data"]:
                    time.append(a)
                barresult["xaxis"]=time
                mydata={}
                for a in myjson["series"]:
                    mydata[a["name"]]=a["data"]
                barresult["data"]=mydata
                outresult.append(barresult)
            if type=="折线":
                barresult={}
                barresult["type"]="折线"
                barresult["title"]=mytitle
                time=[]
                for a in myjson["xAxis"][0]["data"]:
                    time.append(a)
                barresult["xaxis"]=time
                mydata={}
                for a in myjson["series"]:
                    updata=[]
                    for b in a["data"]:
                        updata.append(b[1])
                    mydata[a["name"]]=updata
                barresult["data"]=mydata
                outresult.append(barresult)
            if mytitle=="全部规则相关点各类型占比" or mytitle=="全部规则相关点中新颖性风险点占比":
                barresult={}
                barresult["type"]="饼状"
                barresult["title"]=mytitle
                time=[]
                mydata=[]
                a = myjson["series"][0]["data"]
                for b in a:
                    mydata.append([b["name"],b["value"]])

                barresult["data"]=mydata
                outresult.append(barresult)
        return outresult



    def pdfcreate(self):#生成最后的PDF
        generate_pdf(self.path+str(self.name)+".docx", self.path)
        return self.path+str(self.name)+".pdf"



def generate_pdf(doc_path, path):

    subprocess.call(['soffice', '--convert-to', 'pdf', '--outdir', path, doc_path])
    return path


def myinput(list): #将系统的输入转换成sql语句中所要的格式
    mystr='('
    for i in list:
        mystr+=str(i)
        mystr+=','
    mystr=mystr[0:-1]
    mystr+=')'
    return mystr


def trendsql(stime,etime,list):#趋势分析sql语句
    sqlstr="SELECT substring(main_class_list,1,INSTR(main_class_list,'/')-1) as clss_type, substring(publication_date,1,4) as pub_year,count(patent_type) as cnt FROM patent.patent where substring(publication_date,1,4)<="+str(etime)+" "+"and substring(publication_date,1,4)>="+str(stime)+" "+"and id in"+myinput(list)+" "+"group by substring(main_class_list,1,INSTR(main_class_list,'/')-1),pub_year"
    return sqlstr


def authorsql(list):#发明人分析sql语句
    sqlstr="select id,inventor_list from patent.patent where id in "+myinput(list)
    return sqlstr


def areasql(list):#地域分析sql语句
    sqlstr="select count(c2),c2 from(SELECT  applicant_area,CONCAT(substring(applicant_area,1,INSTR(applicant_area,';')-1) ,substring(applicant_area,1,INSTR(applicant_area,'(')-1))  as c2 FROM patent.patent where id in "+myinput(list)+" ) as ttt group by c2"
    return sqlstr


def timecal(data):#时间置信区间计算
    num=[]
    for i in data:
        num.append(int(i[1]))
    result=st.norm.interval(alpha=0.99,loc=np.mean(num),scale=st.sem(num))
    output=[]
    if int(result[0]) == int(result[1]):
        output.append(int(result[0])-1)
        output.append(int(result[1])+1)
    else:
        output.append(int(result[0]))
        output.append(int(result[1]))
    return output


def timesql(list):#置信区间计算SQL语句
    sqlstr="SELECT  id,substring(publication_date,1,4) as pub_year FROM patent.patent where id in "+myinput(list)
    return sqlstr

# #申请人分析
def author(list,type):
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    output=[]
    sqls=authorsql(list)
    db.myexecu(sqls)
    dd=authoranaly(db,'234')
    dic=dd.authorlist()
    dic=dd.listsort(dic,myinput(list))
    if type == "bar":
       output.append(dd.authorbar(dic))
       db.endconn()
       return output
    output.append(dd.authorpie(dic,'234'))
    db.endconn()
    return output

def pic1(id):
    c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
    xlist=[]
    a=0
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    db.myexecu("SELECT relevant_sig,word_pairs, trigger_rules,index_num FROM patent.novelty_ana_item where novelty_ana_id="+str(id))
    for i in db.data:
        xlist.append("权利"+str(i[3]))
    Y=[]
    c.add_xaxis(xlist)
    for i in db.data:
        Y.append(i[1])
    c.add_yaxis( series_name="相关词",y_axis=Y)
    Yy=[]
    for i in db.data:
        Yy.append(i[2])
    c.add_yaxis( series_name="规则相关点",y_axis=Yy)
    c.set_global_opts(title_opts=opts.TitleOpts(title="各比对结果相关词和规则相关点数量",pos_left="center", pos_top="top"),legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
    return c.dump_options_with_quotes()


def pic2(id):
    c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
    xlist=[]
    a=0
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    db.myexecu("SELECT  relevant_sig,direct_substitution,hyponym_hypernym,numeric_range,index_num FROM patent.novelty_ana_item where novelty_ana_id="+str(id))
    for i in db.data:
        xlist.append("权利"+str(i[4]))
    Y=[]
    c.add_xaxis(xlist)
    for i in db.data:
        Y.append(i[1])
    c.add_yaxis( series_name="概念直接替换",y_axis=Y)
    Yy=[]
    for i in db.data:
        Yy.append(i[2])
    c.add_yaxis( series_name="上下位关系",y_axis=Yy)
    Xx=[]
    for i in db.data:
        Xx.append(i[3])
    c.add_yaxis( series_name="数字范围",y_axis=Xx)
    c.set_global_opts(title_opts=opts.TitleOpts(title="规则相关点各类型数量",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
    return c.dump_options_with_quotes()

def pic3(id):
    c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
    xlist=[]
    a=0
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    db.myexecu("SELECT  relevant_sig,destroy,index_num FROM patent.novelty_ana_item where novelty_ana_id="+str(id))
    for i in db.data:
        xlist.append("权利"+str(i[2]))
    Y=[]
    c.add_xaxis(xlist)
    for i in db.data:
        Y.append(i[1])
    c.add_yaxis( series_name="新颖性风险点",y_axis=Y)
    c.set_global_opts(title_opts=opts.TitleOpts(title="各比对结果疑似新颖性风险点",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
    return c.dump_options_with_quotes()

def pic4(id):
    c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    db.myexecu("SELECT sum(direct_substitution),sum( hyponym_hypernym),sum( numeric_range) FROM patent.novelty_ana_item where novelty_ana_id="+str(id))
    result=[]
    result.append(["概念直接替换",db.data[0][0]])
    result.append(["上下位关系",db.data[0][1]])
    result.append(["数字范围",db.data[0][2]])
    c.add("",result)
    c.set_global_opts(title_opts=opts.TitleOpts(title="全部规则相关点中各类型占比",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
    c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
    return c.dump_options_with_quotes()

def pic5(id):
    c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    db.myexecu("SELECT sum(destroy),sum(trigger_rules) FROM patent.novelty_ana_item where novelty_ana_id="+str(id))
    result=[]
    result.append(["新颖性风险点",db.data[0][0]])
    result.append(["非风险点",db.data[0][1]-db.data[0][0]])
    c.add("",result)
    c.set_global_opts(title_opts=opts.TitleOpts(title="全部规则相关点中新颖性风险点占比",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
    c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
    return c.dump_options_with_quotes()

def novelty_stats_draw(id):
    output=[]
    output.append(pic1(id))
    output.append(pic2(id))
    output.append(pic3(id))
    output.append(pic4(id))
    output.append(pic5(id))
    return output

# #趋势分析
def trend(list,type):
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    sqls=timesql(list)
    db.myexecu(sqls)
    timeblock=timecal(db.data)
    output=[]
    sqls=trendsql(timeblock[0],timeblock[1],list)
    db.myexecu(sqls)
    dd=trendanaly(db,"234")
    dic=dd.newdic(timeblock[0],timeblock[1])
    if type == "bar":
        output.append(dd.trendbar(dic,timeblock[0],timeblock[1]))
        return output
    if type == "line":
        output.append(dd.trendline(dic,timeblock[0],timeblock[1]))
        return output
    for i in dd.trendpie(dic):
        output.append(i)
    db.endconn()
    return output

# #地域分析
def area(list,type):
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    output=[]
    sqls=areasql(list)
    db.myexecu(sqls)
    dd=areaanaly(db,'234')
    dic=dd.newdic()
    if type == "pie":
        output.append(dd.areapie(dic))
        db.endconn()
        return output
    output.append(dd.areabar(dic))
    db.endconn()
    return output

def analyze_by_list(patentIds, figType, anaType):#三个功能封装在一起
    if anaType == 'author':
        return author(patentIds,figType)
    if anaType == 'area':
        return area(patentIds,figType)
    return trend(patentIds,figType)

def pdf_output(id):#输出pdf
    db = mydb('10.108.119.71','zym','zym','patent',3306,'utf8')
    mypdf=pdfcreator(id,db)
    mypdf.finalword()
    output=mypdf.pdfcreate()
    db.myexecu('update patent.report2generate set status="已生成",pdf_file_path="'+output+'" where report_id='+str(mypdf.id))
    db.endconn()
    print("报告生成完毕")
    return

















